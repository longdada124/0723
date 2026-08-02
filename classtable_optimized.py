import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from io import BytesIO
import re
import os
import json
import base64
import unicodedata

try:
    import requests
    _REQUESTS_AVAILABLE = True
except ImportError:
    _REQUESTS_AVAILABLE = False

st.set_page_config(page_title="課表彙整系統", layout="wide")

# ============================================================
# 常數設定
# ============================================================
DAYS = range(1, 6)      # 週一 ~ 週五
PERIODS = range(1, 9)   # 第1節 ~ 第8節

# 計算教師「教學總時數」統計時要排除的節次（例如第8節通常是課後輔導/社團，不計入正式教學時數）。
# 課表本身仍會照常顯示這些節次的排課內容，只有統計數字（BASE/TOTAL/EXTRA）不列入計算。
HOURS_EXCLUDED_PERIODS = {8}

TEMPLATE_FILES = {
    "class": "班級樣板.docx",
    "teacher": "教師樣板.docx",
}

DOWNLOAD_TEMPLATES = {
    "1. 配課表範本": "配課表.xlsx",
    "2. 課表範本": "課表.xlsx",
    "3. 教師排序表範本": "教師排序表.xlsx",
}

GITHUB_API_BASE = "https://api.github.com"

# 支援多種常見的星期表示法（全形數字未涵蓋，如有需要可再擴充）
DAY_MAP = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "週一": 1, "週二": 2, "週三": 3, "週四": 4, "週五": 5,
    "周一": 1, "周二": 2, "周三": 3, "周四": 4, "周五": 5,
    "1": 1, "2": 2, "3": 3, "4": 4, "5": 5,
    "mon": 1, "tue": 2, "wed": 3, "thu": 4, "fri": 5,
}


class DataValidationError(Exception):
    """
    資料格式或內容有問題時拋出。會被主流程攔截並顯示成使用者看得懂的訊息，
    而不是讓 Streamlit 顯示整串 Python 錯誤堆疊。
    """
    pass


# ============================================================
# 共用工具函式
# ============================================================
def normalize_name(name) -> str:
    """去除頭尾空白與全形空白，避免同一位教師因空白差異被誤判成不同人。"""
    if name is None:
        return ""
    return str(name).strip().replace("　", "").replace(" ", "")


def smart_read_csv(uploaded_file):
    """
    依序嘗試常見編碼讀取 CSV。
    台灣 Excel 匯出的 CSV 常見為 Big5／UTF-8-BOM，預設 utf-8 常會直接噴錯，
    這裡自動輪流嘗試，減少使用者因編碼問題而卡關。
    """
    encodings = ["utf-8-sig", "utf-8", "cp950", "big5"]
    last_err = None
    for enc in encodings:
        try:
            uploaded_file.seek(0)
            return pd.read_csv(uploaded_file, encoding=enc)
        except (UnicodeDecodeError, UnicodeError) as e:
            last_err = e
            continue
    raise DataValidationError(
        f"無法辨識 CSV 檔案編碼（已嘗試 UTF-8 / Big5 等格式）。"
        f"建議用 Excel 開啟後另存新檔為 .xlsx 再上傳。\n技術訊息：{last_err}"
    )


def read_uploaded_table(uploaded_file, file_label):
    """統一讀取上傳的 Excel/CSV，失敗時拋出好懂的錯誤訊息並清理欄名空白。"""
    try:
        if uploaded_file.name.lower().endswith(".csv"):
            df = smart_read_csv(uploaded_file)
        else:
            df = pd.read_excel(uploaded_file)
    except DataValidationError:
        raise
    except Exception as e:
        raise DataValidationError(
            f"讀取「{file_label}」失敗，請確認檔案未損毀且格式正確。\n技術訊息：{e}"
        )

    df.columns = [str(c).strip() for c in df.columns]
    if df.empty:
        raise DataValidationError(f"「{file_label}」讀取成功，但內容是空的，請確認檔案內有資料。")
    return df


def validate_columns(df, required_cols, file_label):
    """檢查必要欄位是否存在，缺少時拋出清楚列出缺少哪些欄位的錯誤。"""
    missing = [c for c in required_cols if c not in df.columns]
    if missing:
        raise DataValidationError(
            f"「{file_label}」缺少必要欄位：{'、'.join(missing)}\n"
            f"目前偵測到的欄位為：{'、'.join(str(c) for c in df.columns)}"
        )


def load_default_template(file_name):
    """讀取內建樣板（編碼相容版），支援 NFC/NFD 檔名正規化比對。"""
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        full_path = os.path.join(current_dir, file_name)

        if os.path.exists(full_path):
            with open(full_path, "rb") as f:
                return f.read()

        target_filename_nfc = unicodedata.normalize("NFC", file_name)
        for actual_file in os.listdir(current_dir):
            actual_file_nfc = unicodedata.normalize("NFC", actual_file)
            if actual_file_nfc == target_filename_nfc:
                with open(os.path.join(current_dir, actual_file), "rb") as f:
                    return f.read()
    except Exception:
        pass
    return None


def apply_replacements(doc_obj, replacements: dict, color_map: dict = None):
    """
    一次套用一整組佔位符取代，取代舊版「每個佔位符各自呼叫一次、
    每次都重新掃描整份文件」的做法。

    效能重點：
      1. 段落＋表格清單只蒐集一次（舊版每呼叫一次 master_replace 就重蒐集一次）。
      2. 對每個段落先做一次「{{」快速排除，沒有佔位符的段落直接跳過，
         不必逐一比對 80 幾個佔位符字串。
      3. 對合併大量班級／教師課表的批次匯出來說，效果最明顯。

    color_map（選填）：{佔位符: 6碼色碼（不含#）}，若該段落含有對應的佔位符，
    取代完文字後會把整段文字設成指定顏色，用於兼課類型的顏色標記。
    """
    targets = list(doc_obj.paragraphs)
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                targets.extend(cell.paragraphs)

    color_map = color_map or {}
    formatted = {}
    for old_text, new_text in replacements.items():
        if isinstance(new_text, (float, int)):
            formatted[old_text] = str(int(new_text))
        else:
            formatted[old_text] = str(new_text) if (new_text and str(new_text).strip() != "") else ""

    for p in targets:
        full_text = "".join(run.text for run in p.runs)
        if "{{" not in full_text:
            continue
        updated_text = full_text
        matched_color = None
        for old_text, new_val in formatted.items():
            if old_text in updated_text:
                updated_text = updated_text.replace(old_text, new_val)
                if old_text in color_map:
                    matched_color = color_map[old_text]
        if updated_text != full_text:
            for i, run in enumerate(p.runs):
                run.text = updated_text if i == 0 else ""
            if matched_color and p.runs:
                try:
                    p.runs[0].font.color.rgb = RGBColor.from_string(matched_color)
                    p.runs[0].font.bold = True
                except ValueError:
                    pass  # 色碼格式不正確時直接忽略上色/加粗，不影響文字本身的取代結果


def build_class_replacements(cname, class_data, tutors_map):
    """組出單一班級文件所需的所有佔位符取代字典。"""
    repl = {
        "{{CLASS}}": cname,
        "{{TUTOR}}": tutors_map.get(cname, "未設定"),
    }
    for d in DAYS:
        for p in PERIODS:
            v = class_data.get(cname, {}).get((d, p), {"subj": "", "teacher": ""})
            repl[f"{{{{SD{d}P{p}}}}}"] = v["subj"]
            repl[f"{{{{TD{d}P{p}}}}}"] = v["teacher"]
    return repl


def build_teacher_replacements(tname, teacher_data, base_hours, total_counts):
    """組出單一教師文件所需的所有佔位符取代字典。"""
    base = int(base_hours.get(tname, 0))
    total = int(total_counts.get(tname, 0))
    repl = {
        "{{TEACHER}}": tname,
        "{{BASE}}": base,
        "{{TOTAL}}": total,
        "{{EXTRA}}": total - base,
    }
    for d in DAYS:
        for p in PERIODS:
            v = teacher_data.get(tname, {}).get((d, p), {"subj": "", "class": ""})
            repl[f"{{{{CD{d}P{p}}}}}"] = v["class"]
            repl[f"{{{{SD{d}P{p}}}}}"] = v["subj"]
    return repl


def build_teacher_color_map(tname, duty_marks, duty_types):
    """
    依照兼課標記，組出教師課表匯出 Word 時每個節次需要上色的佔位符對照表。
    同時上色 CD（班級）與 SD（科目）兩個佔位符，確保不管樣板怎麼排版，
    整格內容都會套用兼課類型對應的顏色。
    """
    marks = duty_marks.get(tname, {})
    color_map = {}
    for (d, p), duty_name in marks.items():
        color = duty_types.get(duty_name)
        if color:
            hex_color = str(color).lstrip("#")
            color_map[f"{{{{CD{d}P{p}}}}}"] = hex_color
            color_map[f"{{{{SD{d}P{p}}}}}"] = hex_color
    return color_map


def generate_document(template_bytes, replacements, color_map=None):
    doc = Document(BytesIO(template_bytes))
    apply_replacements(doc, replacements, color_map=color_map)
    return doc


def doc_to_bytes(doc):
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# 資料解析
# ============================================================
_PANDAS_DEDUP_SUFFIX_PATTERN = re.compile(r"\.\d+$")


def _strip_pandas_dedup_suffix(name: str) -> str:
    """
    Excel 裡如果有兩個欄位剛好同名（例如班級ID欄跟某個科目欄剛好都叫「班級」），
    pandas 讀取時會自動把後面重複出現的欄位改名成「班級.1」「班級.2」來避免衝突。
    這裡把這個自動加上去的尾碼還原掉，這樣科目名稱才能跟課表正確比對，
    不會因為欄位剛好跟保留欄名（班級／導師）撞名，就找不到對應教師。
    """
    return _PANDAS_DEDUP_SUFFIX_PATTERN.sub("", str(name))


def parse_assignment(df_assign):
    """
    解析配課表，支援兩種格式：
      模式 A：一維清單格式（需有「科目」「教師」欄位）
      模式 B：二維矩陣格式（科目當欄位名稱，儲存格內容為授課教師）

    回傳：
      assign_lookup: 明細列表 [{'c':班級,'s':科目,'t':教師}, ...]
      assign_dict:   {(班級,科目): [教師,...]}，供課表比對時 O(1) 查詢（效能優化）
      tutors:        {班級: 導師}
      all_teachers_db: 教師名稱集合
      warnings:      解析過程中的提醒訊息
    """
    warnings = []
    validate_columns(df_assign, ["班級"], "配課表")

    assign_lookup = []
    assign_dict = {}
    all_teachers_db = set()
    tutors = {}

    def register(c, s, t_raw):
        t_list = [normalize_name(x) for x in str(t_raw).split("/")]
        valid = [t for t in t_list if t and t != "nan"]
        for t in valid:
            assign_lookup.append({"c": c, "s": s, "t": t})
            all_teachers_db.add(t)
        # 無論教師欄位是否填寫，都建立這組 (班級,科目) 的鍵值。
        # 這樣課表比對時才能區分「配課表已登記但教師欄位留空」（顯示空白即可）
        # 與「配課表中根本沒有這個班級/科目組合」（顯示未知教師並提出警告）兩種情況。
        assign_dict.setdefault((c, s), [])
        assign_dict[(c, s)].extend(valid)
        return valid

    if "科目" in df_assign.columns and "教師" in df_assign.columns:
        # 【模式 A】傳統一維清單格式
        for _, row in df_assign.iterrows():
            c = str(row["班級"]).strip()
            s = str(row["科目"]).strip()
            t_raw = str(row.get("教師", "")).strip()
            register(c, s, t_raw)
            if s == "班級":
                tutors[c] = t_raw
    else:
        # 【模式 B】二維矩陣表格格式
        if "導師" in df_assign.columns:
            for _, row in df_assign.iterrows():
                c = str(row["班級"]).strip()
                t_tutor = str(row["導師"]).strip()
                if t_tutor and t_tutor != "nan":
                    tutors[c] = t_tutor

        subject_cols = [col for col in df_assign.columns if col not in ["班級", "導師"]]
        if not subject_cols:
            raise DataValidationError("「配課表」除了班級／導師欄位外，找不到任何科目欄位，請確認檔案格式。")

        df_melted = pd.melt(df_assign, id_vars=["班級"], value_vars=subject_cols,
                             var_name="科目", value_name="教師")
        for _, row in df_melted.iterrows():
            c = str(row["班級"]).strip()
            # 還原可能被 pandas 自動加上的重複欄位尾碼（見 _strip_pandas_dedup_suffix 說明）
            s = _strip_pandas_dedup_suffix(str(row["科目"]).strip())
            t_raw = str(row["教師"]).strip() if pd.notna(row["教師"]) else ""
            register(c, s, t_raw)

    if not assign_lookup:
        raise DataValidationError("配課表中沒有解析出任何有效的教師配課資料，請確認欄位內容是否正確。")

    return assign_lookup, assign_dict, tutors, all_teachers_db, warnings


def parse_teacher_sort(df_sort, all_teachers_list):
    """
    解析教師排序暨時數表。若使用者未上傳，則以配課表中出現的教師依字母排序、
    應授時數預設為 0。教師姓名比對時會忽略頭尾／全形空白差異。
    """
    warnings = []
    ordered_teachers, base_hours = [], {}

    if df_sort is not None:
        if df_sort.shape[1] < 2:
            raise DataValidationError("「教師排序暨時數表」至少需要兩欄（教師姓名、應授時數）。")

        norm_to_actual = {normalize_name(t): t for t in all_teachers_list}
        for _, s_row in df_sort.iterrows():
            raw_name = str(s_row.iloc[0]).strip()
            key = normalize_name(raw_name)
            if key in norm_to_actual:
                t_name = norm_to_actual[key]
                if t_name not in ordered_teachers:
                    ordered_teachers.append(t_name)
                    try:
                        base_hours[t_name] = int(s_row.iloc[1])
                    except Exception:
                        base_hours[t_name] = 0
            elif raw_name and raw_name != "nan":
                warnings.append(f"教師排序表中的「{raw_name}」在配課表中查無對應資料，已略過。")

        for t in all_teachers_list:
            if t not in ordered_teachers:
                ordered_teachers.append(t)
                base_hours[t] = 0
    else:
        ordered_teachers = sorted(all_teachers_list)
        base_hours = {t: 0 for t in ordered_teachers}

    return ordered_teachers, base_hours, warnings


def _register_schedule_entry(class_data, teacher_data, total_counts, warnings, c_raw, s_raw, d, p, assign_dict):
    """登記單一節課的資料，供一維／矩陣兩種課表格式共用，避免重複邏輯。"""
    if not s_raw or s_raw == "nan":
        display_t, s_raw, curr_t_list = "", "", []
    elif (c_raw, s_raw) in assign_dict:
        # 配課表中有登記這個班級/科目組合。若教師欄位當初留空，這裡就是空清單，
        # 顯示空白即可（例如社團活動、班會等不一定要填教師的科目），不需要警告。
        curr_t_list = assign_dict[(c_raw, s_raw)]
        display_t = "/".join(curr_t_list)
    else:
        # 配課表中完全查無這個班級/科目組合，較可能是資料填寫錯誤或科目名稱打錯字，才需要提醒。
        curr_t_list = []
        display_t = "未知教師"
        warnings.append(f"{c_raw} 週{d}第{p}節「{s_raw}」在配課表中查無對應教師。")

    class_data.setdefault(c_raw, {})[(d, p)] = {"subj": s_raw, "teacher": display_t}
    for t in curr_t_list:
        teacher_data.setdefault(t, {})[(d, p)] = {"subj": s_raw, "class": c_raw}
        if p not in HOURS_EXCLUDED_PERIODS:
            total_counts[t] = total_counts.get(t, 0) + 1


def parse_schedule(df_time, assign_dict):
    """
    解析【一維清單格式】課表（需有「班級」「科目」「星期」「節次」欄位，每列代表一節課）。
    使用字典查詢取代舊版逐筆掃描 assign_lookup 的線性比對，是效能優化的重點之一。
    """
    warnings = []
    validate_columns(df_time, ["班級", "科目", "星期", "節次"], "課表")

    class_data, teacher_data, total_counts = {}, {}, {}
    unmatched_format = 0

    for _, row in df_time.iterrows():
        c_raw = str(row["班級"]).strip()
        s_raw = str(row["科目"]).strip()
        day_key = str(row["星期"]).strip().lower()
        d = DAY_MAP.get(day_key, 0)
        p_match = re.search(r"\d+", str(row["節次"]))

        if not (p_match and d > 0):
            unmatched_format += 1
            continue
        p = int(p_match.group())
        if p not in PERIODS:
            continue

        _register_schedule_entry(class_data, teacher_data, total_counts, warnings, c_raw, s_raw, d, p, assign_dict)

    if unmatched_format:
        warnings.append(f"共有 {unmatched_format} 筆課表資料的「星期」或「節次」格式無法辨識，已自動略過。")

    if not class_data:
        raise DataValidationError("課表中沒有解析出任何有效的班級課程資料，請確認欄位內容與格式是否正確。")

    return class_data, teacher_data, total_counts, warnings


# 矩陣型課表的列標籤格式，例如「一-1」「週二-3」，允許常見的破折號變體
_SCHEDULE_ROW_LABEL_PATTERN = re.compile(r"^(.+?)[-_－﹣]\s*(\d+)$")


def parse_schedule_matrix(df_matrix, assign_dict):
    """
    解析【矩陣格式】課表：
      - 第一欄（表頭通常留白）為「星期-節次」，例如「一-1」「二-3」
      - 其餘欄位表頭為班級名稱（例如 701、702…）
      - 儲存格內容為該節次的科目名稱，空白代表沒有排課
    會自動把矩陣攤平成與一維格式相同的 class_data / teacher_data 結構。
    """
    warnings = []
    if df_matrix.shape[1] < 2:
        raise DataValidationError("「課表」欄位數量不足，無法辨識矩陣格式（需至少一欄星期節次＋一欄以上的班級）。")

    row_label_col = df_matrix.columns[0]
    class_cols = list(df_matrix.columns[1:])

    class_data, teacher_data, total_counts = {}, {}, {}
    unmatched_format = 0

    for _, row in df_matrix.iterrows():
        label = str(row[row_label_col]).strip()
        m = _SCHEDULE_ROW_LABEL_PATTERN.match(label)
        if not m:
            unmatched_format += 1
            continue

        day_key = m.group(1).strip().lower()
        d = DAY_MAP.get(day_key, 0)
        p = int(m.group(2))
        if d == 0 or p not in PERIODS:
            unmatched_format += 1
            continue

        for c_col in class_cols:
            c_raw = str(c_col).strip()
            s_val = row[c_col]
            s_raw = "" if pd.isna(s_val) else str(s_val).strip()
            _register_schedule_entry(class_data, teacher_data, total_counts, warnings, c_raw, s_raw, d, p, assign_dict)

    if unmatched_format:
        warnings.append(f"共有 {unmatched_format} 列「星期-節次」格式無法辨識（例如應為「一-1」），已自動略過。")

    if not class_data:
        raise DataValidationError("課表（矩陣格式）中沒有解析出任何有效的班級課程資料，請確認格式是否正確。")

    return class_data, teacher_data, total_counts, warnings


def parse_schedule_auto(df_time, assign_dict):
    """
    自動辨識課表檔案是哪一種格式並解析：
      - 一維清單格式：欄位包含「班級」「科目」「星期」「節次」
      - 矩陣格式：第一欄為星期-節次（如「一-1」），其餘欄位表頭為班級名稱
    """
    long_format_cols = {"班級", "科目", "星期", "節次"}
    if long_format_cols.issubset(set(df_time.columns)):
        return parse_schedule(df_time, assign_dict)
    return parse_schedule_matrix(df_time, assign_dict)


def collect_class_subject_combos(class_data, assign_dict=None):
    """
    取得節數檢查需要比對的所有 (班級,科目) 組合，取兩個來源的聯集：
      1. 實際課表中出現過的科目（不含空堂）
      2. 配課表中登記過的科目（不含用來標記導師的特殊列「班級」）
    如果只看第 1 項，會漏掉「配課表配了老師、但整學期完全沒被排進課表」的科目，
    第 2 項就是專門補上這個漏洞。
    """
    combos = set()
    for c, periods in class_data.items():
        for info in periods.values():
            subj = info.get("subj", "")
            if subj:
                combos.add((c, subj))
    if assign_dict:
        for (c, s) in assign_dict.keys():
            if s and s != "班級":
                combos.add((c, s))
    return combos


def compute_actual_subject_counts(class_data):
    """統計目前課表中，每個 (班級,科目) 實際被排了幾節課（不含空堂）。"""
    counts = {}
    for c, periods in class_data.items():
        for info in periods.values():
            subj = info.get("subj", "")
            if subj:
                key = (c, subj)
                counts[key] = counts.get(key, 0) + 1
    return counts


def detect_teacher_conflicts(class_data):
    """
    偵測「同一位教師在同一時段被排到多個不同班級」的排課衝突。
    以 class_data（各班課表）為基礎逐一攤開比對，一節課若由多位教師合開
    （以「/」分隔），會分別檢查每一位教師。
    回傳 list，每筆為 {'teacher','day','period','classes'}，並依 週幾/第幾節 排序。
    """
    slot_map = {}
    for c, periods in class_data.items():
        for (d, p), info in periods.items():
            teacher_str = info.get("teacher", "")
            if not teacher_str:
                continue
            for t in teacher_str.split("/"):
                t = t.strip()
                if not t or t == "未知教師":
                    continue
                slot_map.setdefault((t, d, p), set()).add(c)

    conflicts = [
        {"teacher": t, "day": d, "period": p, "classes": sorted(classes)}
        for (t, d, p), classes in slot_map.items()
        if len(classes) > 1
    ]
    conflicts.sort(key=lambda x: (x["day"], x["period"], x["teacher"]))
    return conflicts


def resolve_teacher_display(c_raw, s_raw, assign_dict):
    """
    依照配課表判斷某節課應顯示的教師文字，邏輯與匯入課表時的 _register_schedule_entry 一致，
    供「網頁上直接編輯課表」功能使用，確保手動編輯跟匯入的結果呈現一致。
    """
    s_raw = (s_raw or "").strip()
    if not s_raw or s_raw == "nan":
        return "", ""
    if (c_raw, s_raw) in assign_dict:
        curr_t_list = assign_dict[(c_raw, s_raw)]
        return s_raw, ("/".join(curr_t_list) if curr_t_list else "")
    return s_raw, "未知教師"


def rebuild_teacher_data(class_data):
    """
    根據目前的 class_data（各班課表，唯一的真實來源）重新計算 teacher_data 與 total_counts。
    用於使用者在網頁上手動編輯某個班級的課表之後，確保教師課表與統計數字保持同步，
    避免手動增修造成資料不一致。
    """
    teacher_data, total_counts = {}, {}
    for c, periods in class_data.items():
        for (d, p), info in periods.items():
            teacher_str = info.get("teacher", "")
            if not teacher_str:
                continue
            for t in teacher_str.split("/"):
                t = t.strip()
                if not t or t == "未知教師":
                    continue
                teacher_data.setdefault(t, {})[(d, p)] = {"subj": info.get("subj", ""), "class": c}
                if p not in HOURS_EXCLUDED_PERIODS:
                    total_counts[t] = total_counts.get(t, 0) + 1
    return teacher_data, total_counts


# ============================================================
# 雲端持久化：把整合後的資料存回 GitHub Repo，下次打開網頁自動還原
# ============================================================
def get_github_config():
    """
    從 st.secrets 讀取 GitHub 存檔設定。
    需在 Streamlit Cloud 的 App settings → Secrets 加入類似：
        [github]
        token = "ghp_xxxxxxxxxxxxxxxxxxxx"
        repo = "your-username/your-repo-name"
        branch = "main"
        data_path = "saved_data.json"
    若未設定，回傳 None，代表雲端存檔功能停用（不影響其他功能正常使用）。
    """
    if not _REQUESTS_AVAILABLE:
        return None
    try:
        gh = st.secrets["github"]
        return {
            "token": gh["token"],
            "repo": gh["repo"],
            "branch": gh.get("branch", "main"),
            "path": gh.get("data_path", "saved_data.json"),
        }
    except Exception:
        return None


def serialize_period_dict(data: dict) -> dict:
    """
    class_data / teacher_data 的 key 是 (星期, 節次) tuple，JSON 不支援 tuple 當 key，
    這裡轉成 "d_p" 字串格式方便存成 JSON，讀回時再還原成 tuple。
    """
    return {
        name: {f"{d}_{p}": info for (d, p), info in periods.items()}
        for name, periods in data.items()
    }


def deserialize_period_dict(data: dict) -> dict:
    out = {}
    for name, periods in data.items():
        out[name] = {}
        for key, info in periods.items():
            d_str, p_str = key.split("_")
            out[name][(int(d_str), int(p_str))] = info
    return out


def build_save_payload():
    """把目前 session_state 中的整合結果打包成可存成 JSON 的字典。"""
    return {
        "class_data": serialize_period_dict(st.session_state.class_data),
        "teacher_data": serialize_period_dict(st.session_state.teacher_data),
        "tutors_map": st.session_state.tutors_map,
        "base_hours": st.session_state.base_hours,
        "total_counts": st.session_state.total_counts,
        "ordered_teachers": st.session_state.ordered_teachers,
        "df_assign": st.session_state.df_assign.to_dict(orient="records"),
        "import_warnings": st.session_state.get("import_warnings", []),
        "expected_hours": st.session_state.get("expected_hours", {}),
        "duty_types": st.session_state.get("duty_types", {}),
        "duty_marks": serialize_period_dict(st.session_state.get("duty_marks", {})),
    }


def restore_from_payload(payload: dict):
    """把從雲端讀回的 JSON 資料還原成 session_state 需要的格式，並寫入 session_state。"""
    class_data = deserialize_period_dict(payload["class_data"])
    teacher_data = deserialize_period_dict(payload["teacher_data"])
    df_assign = pd.DataFrame(payload["df_assign"])
    ordered_teachers = payload.get("ordered_teachers", [])

    st.session_state.update({
        "class_data": class_data,
        "teacher_data": teacher_data,
        "tutors_map": payload.get("tutors_map", {}),
        "base_hours": payload.get("base_hours", {}),
        "total_counts": payload.get("total_counts", {}),
        "ordered_teachers": ordered_teachers,
        "df_assign": df_assign,
        "import_warnings": payload.get("import_warnings", []),
        "expected_hours": payload.get("expected_hours", {}),
        "duty_types": payload.get("duty_types", {}),
        "duty_marks": deserialize_period_dict(payload.get("duty_marks", {})),
        "sel_class": sorted(class_data.keys())[0] if class_data else None,
        "sel_teacher": ordered_teachers[0] if ordered_teachers else None,
        "class_template": load_default_template(TEMPLATE_FILES["class"]),
        "teacher_template": load_default_template(TEMPLATE_FILES["teacher"]),
    })


def _format_github_http_error(status_code, text, action="操作"):
    """把 GitHub API 常見錯誤碼轉成使用者看得懂、附排查建議的訊息。"""
    if status_code == 403 and "not accessible" in text.lower():
        return (
            f"雲端{action}失敗（HTTP 403：Token 沒有寫入/讀取權限）。請依序檢查：\n"
            "1. Token 的「Repository access」是否確實包含這個 repo\n"
            "2. Token 的「Permissions → Contents」是否設為「Read and write」\n"
            "3. 若 repo 屬於 Organization（組織），Fine-grained token 需要組織管理員核准，"
            "請確認 token 狀態不是「Pending approval」\n"
            "4. secrets 裡的 repo 欄位格式是否為「帳號名稱/repo名稱」，拼字與大小寫是否正確\n"
            "（若一直卡關，改用 Classic Token 並勾選「repo」scope 通常最不容易出錯）\n"
            f"原始錯誤：{text[:200]}"
        )
    if status_code == 401:
        return f"雲端{action}失敗（HTTP 401：Token 無效或已過期），請重新產生一組 Token 並更新 secrets。"
    if status_code == 404:
        return f"雲端{action}失敗（HTTP 404：找不到 repo 或分支），請確認 secrets 裡的 repo／branch 名稱是否正確。"
    return f"雲端{action}失敗（HTTP {status_code}）：{text[:300]}"


def github_save_data(payload: dict):
    """將整合結果寫回 GitHub Repo（檔案不存在就新增，存在就更新）。"""
    cfg = get_github_config()
    if not cfg:
        return False, "尚未設定 GitHub 存檔金鑰，已略過雲端儲存（僅存在目前這次連線的記憶體中）。"

    api_url = f"{GITHUB_API_BASE}/repos/{cfg['repo']}/contents/{cfg['path']}"
    headers = {"Authorization": f"token {cfg['token']}", "Accept": "application/vnd.github+json"}

    sha = None
    try:
        r = requests.get(api_url, headers=headers, params={"ref": cfg["branch"]}, timeout=15)
        if r.status_code == 200:
            sha = r.json().get("sha")
    except requests.RequestException as e:
        return False, f"連線 GitHub 失敗：{e}"

    content_str = json.dumps(payload, ensure_ascii=False, indent=2)
    content_b64 = base64.b64encode(content_str.encode("utf-8")).decode("utf-8")
    body = {"message": "📚 自動更新課表整合資料", "content": content_b64, "branch": cfg["branch"]}
    if sha:
        body["sha"] = sha

    try:
        r = requests.put(api_url, headers=headers, json=body, timeout=15)
        if r.status_code in (200, 201):
            return True, "已成功儲存至 GitHub，下次打開網頁會自動沿用這份資料。"
        return False, _format_github_http_error(r.status_code, r.text, action="儲存")
    except requests.RequestException as e:
        return False, f"連線 GitHub 失敗：{e}"


def github_load_data():
    """從 GitHub Repo 讀取先前的存檔，回傳 (payload, error_message)。找不到檔案不算錯誤。"""
    cfg = get_github_config()
    if not cfg:
        return None, None

    api_url = f"{GITHUB_API_BASE}/repos/{cfg['repo']}/contents/{cfg['path']}"
    headers = {"Authorization": f"token {cfg['token']}", "Accept": "application/vnd.github+json"}
    try:
        r = requests.get(api_url, headers=headers, params={"ref": cfg["branch"]}, timeout=15)
        if r.status_code == 404:
            return None, None
        if r.status_code != 200:
            return None, _format_github_http_error(r.status_code, r.text, action="讀取")
        content_str = base64.b64decode(r.json().get("content", "")).decode("utf-8")
        return json.loads(content_str), None
    except requests.RequestException as e:
        return None, f"連線 GitHub 失敗：{e}"
    except Exception as e:
        return None, f"雲端存檔格式錯誤，可能已毀損：{e}"


def github_delete_data():
    """清空重置時一併刪除雲端存檔，避免下次打開網頁又被自動還原回來。"""
    cfg = get_github_config()
    if not cfg:
        return
    api_url = f"{GITHUB_API_BASE}/repos/{cfg['repo']}/contents/{cfg['path']}"
    headers = {"Authorization": f"token {cfg['token']}", "Accept": "application/vnd.github+json"}
    try:
        r = requests.get(api_url, headers=headers, params={"ref": cfg["branch"]}, timeout=15)
        if r.status_code == 200:
            sha = r.json().get("sha")
            requests.delete(
                api_url, headers=headers,
                json={"message": "🧹 清空課表整合資料", "sha": sha, "branch": cfg["branch"]},
                timeout=15,
            )
    except requests.RequestException:
        pass


# ============================================================
# 開啟網頁時，若尚未有資料且雲端有存檔，自動還原（不需重新上傳三份檔案）
# ============================================================
if "class_data" not in st.session_state and not st.session_state.get("_cloud_restore_attempted", False):
    st.session_state["_cloud_restore_attempted"] = True  # 避免每次互動都重打 API
    _payload, _err = github_load_data()
    if _payload:
        try:
            restore_from_payload(_payload)
            st.session_state["last_import_summary"] = "☁️ 已自動從雲端存檔載入上次整合完成的資料。"
        except Exception as e:
            st.session_state["_cloud_restore_error"] = f"雲端存檔還原失敗，請重新上傳三份檔案。技術訊息：{e}"
    elif _err:
        st.session_state["_cloud_restore_error"] = _err

# ============================================================
# 側邊欄：資料管理
# ============================================================
with st.sidebar:
    cloud_cfg = get_github_config()
    with st.expander("☁️ 雲端存檔狀態", expanded=False):
        if cloud_cfg:
            st.caption(f"已連接：{cloud_cfg['repo']}（{cloud_cfg['branch']} 分支）")
            if st.session_state.get("_cloud_restore_error"):
                st.warning(st.session_state["_cloud_restore_error"])
            if "class_data" in st.session_state:
                if st.button("🔄 手動重新同步到雲端"):
                    ok, msg = github_save_data(build_save_payload())
                    if ok:
                        st.success(msg)
                    else:
                        st.error(msg)
        elif not _REQUESTS_AVAILABLE:
            st.caption(
                "⚠️ 找不到 requests 套件，雲端存檔功能已停用。"
                "請在 requirements.txt 加入一行「requests」後重新部署。"
            )
        else:
            st.caption(
                "尚未設定雲端存檔（st.secrets 缺少 [github] 設定），"
                "目前僅會暫存在本次連線中，重新整理網頁後仍需重新上傳。"
            )

    st.header("⚙️ 資料管理")

    if "confirm_reset" not in st.session_state:
        st.session_state.confirm_reset = False

    if not st.session_state.confirm_reset:
        if st.button("🧹 清空重置系統"):
            st.session_state.confirm_reset = True
            st.rerun()
    else:
        st.warning("確定要清空所有已匯入的資料嗎？此動作無法復原。")
        cc1, cc2 = st.columns(2)
        if cc1.button("✅ 確定清空"):
            github_delete_data()
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
        if cc2.button("↩️ 取消"):
            st.session_state.confirm_reset = False
            st.rerun()

    st.divider()
    st.subheader("📥 範本下載")
    for label, file_name in DOWNLOAD_TEMPLATES.items():
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            f_path = os.path.join(current_dir, file_name)
            if os.path.exists(f_path):
                with open(f_path, "rb") as f:
                    st.download_button(label=label, data=f, file_name=file_name, key=f"dl_{file_name}")
            else:
                st.caption(f"⚠️ 找不到 {file_name}")
        except Exception:
            st.caption(f"⚠️ 讀取 {file_name} 失敗")
    st.divider()

    st.subheader("📤 上傳資料檔")
    f_assign = st.file_uploader("1. 上傳【配課表】", type=["xlsx", "csv"])
    f_time = st.file_uploader(
        "2. 上傳【課表】", type=["xlsx", "csv"],
        help="支援兩種格式：一維清單（欄位：班級／科目／星期／節次）或矩陣格式（第一欄為「一-1」這種星期-節次，其餘欄位為班級），系統會自動判斷。",
    )
    f_sort = st.file_uploader("3. 上傳【教師排序暨時數表】(選填)", type=["xlsx", "csv"])

    if f_assign and f_time and st.button("🚀 執行整合"):
        class_temp = load_default_template(TEMPLATE_FILES["class"])
        teacher_temp = load_default_template(TEMPLATE_FILES["teacher"])

        if not class_temp or not teacher_temp:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            actual_files = os.listdir(current_dir) if os.path.exists(current_dir) else []
            st.error("❌ 系統錯誤：後台找不到「班級樣板.docx」或「教師樣板.docx」，請確認 GitHub 檔案。")
            st.info(f"🔍 雲端伺服器目前實際看到的檔案清單：\n`{actual_files}`")
        else:
            try:
                with st.spinner("同步內建樣板與解析資料中..."):
                    df_assign = read_uploaded_table(f_assign, "配課表")
                    df_time = read_uploaded_table(f_time, "課表")
                    df_sort = read_uploaded_table(f_sort, "教師排序暨時數表") if f_sort else None

                    assign_lookup, assign_dict, tutors, all_teachers_db, w1 = parse_assignment(df_assign)
                    all_teachers_list = list(all_teachers_db)
                    ordered_teachers, base_hours, w2 = parse_teacher_sort(df_sort, all_teachers_list)
                    class_data, teacher_data, total_counts, w3 = parse_schedule_auto(df_time, assign_dict)

                    conflicts = detect_teacher_conflicts(class_data)
                    w4 = [
                        f"🚨 排課衝突：{cf['teacher']} 老師在 週{cf['day']} 第{cf['period']}節 "
                        f"同時被排到 {'、'.join(cf['classes'])}，請確認是否為人工排課錯誤。"
                        for cf in conflicts
                    ]

                    all_warnings = w1 + w2 + w3 + w4

                st.session_state.update({
                    "class_template": class_temp,
                    "teacher_template": teacher_temp,
                    "df_assign": df_assign,
                    "class_data": class_data,
                    "teacher_data": teacher_data,
                    "tutors_map": tutors,
                    "base_hours": base_hours,
                    "total_counts": total_counts,
                    "ordered_teachers": ordered_teachers,
                    "import_warnings": all_warnings,
                    "sel_class": sorted(class_data.keys())[0],
                    "sel_teacher": ordered_teachers[0],
                })
                ok, cloud_msg = github_save_data(build_save_payload())
                cloud_suffix = f"\n\n☁️ {cloud_msg}"
                conflict_suffix = f"\n\n🚨 偵測到 {len(conflicts)} 筆排課衝突，請至下方「⚠️ 匯入警告」查看詳情。" if conflicts else ""
                st.session_state["last_import_summary"] = (
                    f"✅ 整合完成！共解析 {len(class_data)} 個班級、"
                    f"{len(ordered_teachers)} 位教師、{len(assign_lookup)} 筆配課紀錄。"
                    f"{conflict_suffix}{cloud_suffix}"
                )
                st.rerun()
            except DataValidationError as e:
                st.error(f"❌ 資料檢查未通過：\n\n{e}")
            except Exception as e:
                st.error(f"❌ 發生未預期的錯誤，請確認檔案格式是否正確。\n\n技術訊息：{e}")

# ============================================================
# 主介面與預覽
# ============================================================
if "class_data" in st.session_state:

    # 匯入成功摘要（只顯示一次，切換分頁後不再重複出現）
    summary = st.session_state.pop("last_import_summary", None)
    if summary:
        st.success(summary)

    import_warnings = st.session_state.get("import_warnings", [])
    if import_warnings:
        with st.expander(f"⚠️ 匯入時偵測到 {len(import_warnings)} 筆提醒事項，點擊展開查看", expanded=False):
            for w in import_warnings[:50]:
                st.write(f"- {w}")
            if len(import_warnings) > 50:
                st.caption(f"…等共 {len(import_warnings)} 筆，僅顯示前 50 筆。")

    tab1, tab2, tab3, tab4, tab5 = st.tabs(
        ["🏫 班級課表", "👩‍🏫 教師課表", "📋 配課總覽與分頁匯出", "🔍 課表搜尋", "📊 節數檢查"]
    )

    with tab1:
        classes = sorted(st.session_state.class_data.keys())
        curr_c = st.session_state.get("sel_class", classes[0])
        col1, col2, col3 = st.columns([1, 2, 1])
        if col1.button("⬅️ 上一班"):
            st.session_state.sel_class = classes[(classes.index(curr_c) - 1) % len(classes)]
            st.rerun()
        if col3.button("下一班 ➡️"):
            st.session_state.sel_class = classes[(classes.index(curr_c) + 1) % len(classes)]
            st.rerun()
        with col2:
            st.session_state.sel_class = st.selectbox("選取班級", classes, index=classes.index(curr_c))

        target_c = st.session_state.sel_class
        st.info(f"📍 班級：{target_c} ｜ 導師：{st.session_state.tutors_map.get(target_c, '未設定')}")

        c_preview = []
        for p in PERIODS:
            row = {"節次": f"第 {p} 節"}
            for d in DAYS:
                info = st.session_state.class_data[target_c].get((d, p))
                if info:
                    teacher_part = f"\n({info['teacher']})" if info["teacher"] else ""
                    row[f"週{d}"] = f"{info['subj']}{teacher_part}"
                else:
                    row[f"週{d}"] = ""
            c_preview.append(row)
        st.table(pd.DataFrame(c_preview))

        bc1, bc2 = st.columns(2)
        with bc1:
            if st.button(f"📥 下載 {target_c} 課表"):
                repl = build_class_replacements(target_c, st.session_state.class_data, st.session_state.tutors_map)
                doc = generate_document(st.session_state.class_template, repl)
                st.download_button(f"💾 儲存 {target_c} 課表", doc_to_bytes(doc), f"{target_c}_班級課表.docx")
        with bc2:
            sel_c_batch = st.multiselect("勾選批次合併", classes, default=classes)
            if st.button("🚀 執行班級合併列印"):
                main_doc = None
                with st.spinner(f"正在產生 {len(sel_c_batch)} 個班級的課表…"):
                    for i, cname in enumerate(sel_c_batch):
                        repl = build_class_replacements(cname, st.session_state.class_data, st.session_state.tutors_map)
                        tmp = generate_document(st.session_state.class_template, repl)
                        if i == 0:
                            main_doc = tmp
                        else:
                            for el in tmp.element.body:
                                main_doc.element.body.append(el)
                if main_doc:
                    st.download_button("💾 下載班級彙整檔", doc_to_bytes(main_doc), "全校班級課表.docx")

        with st.expander(f"✏️ 修正 {target_c} 的排課錯誤", expanded=False):
            st.caption("直接修改下方表格的科目名稱即可（教師會依照配課表自動帶入，不需手動輸入），改完按下方按鈕儲存。留空代表該節沒有排課。")

            edit_rows = []
            for p in PERIODS:
                row = {"節次": f"第 {p} 節"}
                for d in DAYS:
                    info = st.session_state.class_data[target_c].get((d, p), {})
                    row[f"週{d}"] = info.get("subj", "")
                edit_rows.append(row)
            schedule_edit_df = pd.DataFrame(edit_rows)

            edited_schedule_df = st.data_editor(
                schedule_edit_df,
                use_container_width=True,
                hide_index=True,
                key=f"schedule_editor_{target_c}",
                disabled=["節次"],
            )

            if st.button(f"💾 儲存 {target_c} 的修改", key=f"save_schedule_{target_c}"):
                _al, assign_dict_for_edit, _tu, _at, _w = parse_assignment(st.session_state.df_assign)

                new_periods = {}
                for idx, p in enumerate(PERIODS):
                    for d in DAYS:
                        raw_val = edited_schedule_df.iloc[idx][f"週{d}"]
                        raw_subj = "" if pd.isna(raw_val) else str(raw_val).strip()
                        s_clean, display_t = resolve_teacher_display(target_c, raw_subj, assign_dict_for_edit)
                        new_periods[(d, p)] = {"subj": s_clean, "teacher": display_t}

                st.session_state.class_data[target_c] = new_periods
                new_teacher_data, new_total_counts = rebuild_teacher_data(st.session_state.class_data)
                st.session_state.teacher_data = new_teacher_data
                st.session_state.total_counts = new_total_counts

                ok, msg = github_save_data(build_save_payload())
                if ok:
                    st.success(f"✅ {target_c} 的課表已更新，並已同步至雲端。")
                else:
                    st.warning(f"✅ {target_c} 的課表已更新，但雲端同步失敗：{msg}")
                st.rerun()

    with tab2:
        teachers = st.session_state.ordered_teachers
        curr_t = st.session_state.get("sel_teacher", teachers[0])
        colt1, colt2, colt3 = st.columns([1, 2, 1])
        if colt1.button("⬅️ 前一位"):
            st.session_state.sel_teacher = teachers[(teachers.index(curr_t) - 1) % len(teachers)]
            st.rerun()
        if colt3.button("下一位 ➡️"):
            st.session_state.sel_teacher = teachers[(teachers.index(curr_t) + 1) % len(teachers)]
            st.rerun()
        with colt2:
            st.session_state.sel_teacher = st.selectbox("跳轉教師", teachers, index=teachers.index(curr_t))

        target_t = st.session_state.sel_teacher
        base = int(st.session_state.base_hours.get(target_t, 0))
        total = int(st.session_state.total_counts.get(target_t, 0))
        m1, m2, m3 = st.columns(3)
        m1.metric("應授時數", f"{base} 節")
        m2.metric("教學總時數", f"{total} 節")
        m3.metric("兼代課時數", f"{total - base} 節")

        st.session_state.setdefault("duty_types", {})
        st.session_state.setdefault("duty_marks", {})
        duty_types_cfg = st.session_state.duty_types
        teacher_marks = st.session_state.duty_marks.get(target_t, {})

        t_prev = [
            {
                "節次": f"第 {p} 節",
                **{
                    f"週{d}": f"{st.session_state.teacher_data.get(target_t, {}).get((d, p), {}).get('class', '')} "
                              f"{st.session_state.teacher_data.get(target_t, {}).get((d, p), {}).get('subj', '')}".strip()
                    for d in DAYS
                },
            }
            for p in PERIODS
        ]
        t_prev_df = pd.DataFrame(t_prev)

        if teacher_marks and duty_types_cfg:
            def _highlight_duty_cells(_df):
                styles = pd.DataFrame("", index=_df.index, columns=_df.columns)
                for idx, p in enumerate(PERIODS):
                    for d in DAYS:
                        duty_name = teacher_marks.get((d, p))
                        color = duty_types_cfg.get(duty_name) if duty_name else None
                        if color:
                            styles.iloc[idx, styles.columns.get_loc(f"週{d}")] = f"color: #{color}; font-weight: bold;"
                return styles

            st.dataframe(t_prev_df.style.apply(_highlight_duty_cells, axis=None), use_container_width=True, hide_index=True)
            legend_html = " ｜ ".join(
                f'<span style="color:#{color}; font-weight:bold;">■ {name}</span>'
                for name, color in duty_types_cfg.items()
                if any(teacher_marks.get((d, p)) == name for d in DAYS for p in PERIODS)
            )
            if legend_html:
                st.markdown(f"**圖例**：{legend_html}", unsafe_allow_html=True)
        else:
            st.table(t_prev_df)

        bt1, bt2 = st.columns(2)
        with bt1:
            if st.button(f"📥 下載 {target_t} 課表"):
                repl = build_teacher_replacements(
                    target_t, st.session_state.teacher_data, st.session_state.base_hours, st.session_state.total_counts
                )
                color_map = build_teacher_color_map(target_t, st.session_state.duty_marks, duty_types_cfg)
                doc = generate_document(st.session_state.teacher_template, repl, color_map=color_map)
                st.download_button(f"💾 儲存 {target_t} 課表", doc_to_bytes(doc), f"{target_t}_教師課表.docx")
        with bt2:
            sel_t_batch = st.multiselect("批次合併教師", teachers, default=teachers)
            if st.button("🚀 執行教師合併列印"):
                main_doc = None
                with st.spinner(f"正在產生 {len(sel_t_batch)} 位教師的課表…"):
                    for i, tname in enumerate(sel_t_batch):
                        repl = build_teacher_replacements(
                            tname, st.session_state.teacher_data, st.session_state.base_hours, st.session_state.total_counts
                        )
                        color_map = build_teacher_color_map(tname, st.session_state.duty_marks, duty_types_cfg)
                        tmp = generate_document(st.session_state.teacher_template, repl, color_map=color_map)
                        if i == 0:
                            main_doc = tmp
                        else:
                            for el in tmp.element.body:
                                main_doc.element.body.append(el)
                if main_doc:
                    st.download_button("💾 下載教師彙整檔", doc_to_bytes(main_doc), "全校教師課表_彙整.docx")

        with st.expander("⚙️ 管理兼課種類與顏色", expanded=False):
            st.caption(
                "新增／刪除兼課類型，並設定該類型在課表上要顯示的文字顏色（請輸入6碼色碼，不用加#）。"
                "常見色碼參考：紅色 FF0000／藍色 0000FF／綠色 008000／橘色 FFA500／紫色 800080。"
            )
            dt_rows = [{"類型名稱": k, "顏色色碼": v} for k, v in st.session_state.duty_types.items()]
            dt_df = pd.DataFrame(dt_rows) if dt_rows else pd.DataFrame(columns=["類型名稱", "顏色色碼"])
            edited_dt_df = st.data_editor(
                dt_df, num_rows="dynamic", use_container_width=True, hide_index=True, key="duty_types_editor",
            )

            new_duty_types = {}
            invalid_names = []
            for _, row in edited_dt_df.iterrows():
                name = str(row.get("類型名稱", "") or "").strip()
                color_raw = str(row.get("顏色色碼", "") or "").strip().lstrip("#").upper()
                if not name:
                    continue
                if re.fullmatch(r"[0-9A-F]{6}", color_raw):
                    new_duty_types[name] = color_raw
                else:
                    invalid_names.append(name)
            st.session_state.duty_types = new_duty_types
            if invalid_names:
                st.warning(f"⚠️ 以下類型的顏色色碼格式不正確，已略過：{'、'.join(invalid_names)}（請輸入6碼色碼，例如 FF0000）")

            if st.button("💾 儲存兼課類型設定"):
                ok, msg = github_save_data(build_save_payload())
                if ok:
                    st.success(f"☁️ {msg}")
                else:
                    st.warning(msg)

        with st.expander(f"🖊️ 設定 {target_t} 的兼課標記", expanded=False):
            if not st.session_state.duty_types:
                st.info("請先在上方「⚙️ 管理兼課種類與顏色」新增至少一種兼課類型。")
            else:
                actual_periods = sorted(st.session_state.teacher_data.get(target_t, {}).keys())
                if not actual_periods:
                    st.info(f"{target_t} 目前沒有排課，沒有可以標記的節次。")
                else:
                    st.caption("只會列出這位教師實際有排課的節次，沒有課的節次不會出現，避免誤標。")
                    duty_options = ["（無）"] + list(st.session_state.duty_types.keys())
                    mark_rows = []
                    for (d, p) in actual_periods:
                        info = st.session_state.teacher_data[target_t][(d, p)]
                        cur = teacher_marks.get((d, p), "（無）")
                        mark_rows.append({
                            "星期": f"週{d}", "節次": f"第{p}節",
                            "班級": info.get("class", ""), "科目": info.get("subj", ""),
                            "兼課類型": cur if cur in duty_options else "（無）",
                        })
                    mark_df = pd.DataFrame(mark_rows)

                    edited_mark_df = st.data_editor(
                        mark_df, use_container_width=True, hide_index=True,
                        key=f"duty_marks_editor_{target_t}",
                        disabled=["星期", "節次", "班級", "科目"],
                        column_config={"兼課類型": st.column_config.SelectboxColumn(options=duty_options)},
                    )

                    if st.button(f"💾 儲存 {target_t} 的兼課標記", key=f"save_duty_marks_{target_t}"):
                        new_marks = {}
                        for (d, p), val in zip(actual_periods, edited_mark_df["兼課類型"]):
                            if val and val != "（無）":
                                new_marks[(d, p)] = val
                        st.session_state.duty_marks[target_t] = new_marks
                        ok, msg = github_save_data(build_save_payload())
                        if ok:
                            st.success(f"✅ {target_t} 的兼課標記已更新，並已同步至雲端。")
                        else:
                            st.warning(f"✅ {target_t} 的兼課標記已更新，但雲端同步失敗：{msg}")
                        st.rerun()

    with tab3:
        st.header("📋 全校配課資料總覽")
        if "df_assign" in st.session_state:
            st.write(
                "💡 **提示**：下方顯示您所上傳的原始單一配課表。點擊最下方的按鈕，系統將自動依據"
                "「班級」欄位將資料拆分，產出含有**多個班級分頁**的 Excel 活頁簿。"
            )
            st.dataframe(st.session_state.df_assign, use_container_width=True)
            st.divider()
            st.subheader("📥 匯出「一班一工作表」Excel 檔案")

            buf_excel = BytesIO()
            with pd.ExcelWriter(buf_excel, engine="openpyxl") as writer:
                for cname, group in st.session_state.df_assign.groupby("班級"):
                    clean_sheet_name = str(cname).strip()
                    clean_sheet_name = re.sub(r"[\\/*?:\[\]]", "", clean_sheet_name)[:31]
                    if not clean_sheet_name:
                        clean_sheet_name = "未命名班級"
                    group.to_excel(writer, sheet_name=clean_sheet_name, index=False)

            st.download_button(
                label="💾 點我下載「各班級獨立分頁」配課明細表.xlsx",
                data=buf_excel.getvalue(),
                file_name="全校各班級配課明細表_分頁版.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

    with tab4:
        st.header("🔍 課表關鍵字搜尋")
        st.caption("輸入班級、教師或科目名稱（支援部分關鍵字），快速找出所有符合的排課紀錄。")
        keyword = st.text_input("搜尋關鍵字", placeholder="例如：王小明、302、數學")

        if keyword:
            kw = keyword.strip()
            rows = []
            for c, periods in st.session_state.class_data.items():
                for (d, p), info in periods.items():
                    subj, teacher = info.get("subj", ""), info.get("teacher", "")
                    if kw in c or kw in subj or kw in teacher:
                        rows.append({
                            "班級": c, "_d": d, "_p": p,
                            "星期": f"週{d}", "節次": f"第{p}節",
                            "科目": subj, "教師": teacher,
                        })
            if rows:
                result_df = (
                    pd.DataFrame(rows)
                    .sort_values(["班級", "_d", "_p"])
                    .drop(columns=["_d", "_p"])
                    .reset_index(drop=True)
                )
                st.success(f"共找到 {len(result_df)} 筆符合「{keyword}」的紀錄")
                st.dataframe(result_df, use_container_width=True, hide_index=True)
            else:
                st.warning(f"找不到符合「{keyword}」的紀錄，請確認關鍵字是否正確。")
        else:
            st.info("請輸入關鍵字開始搜尋。")

    with tab5:
        st.header("📊 節數檢查與排課衝突偵測")
        st.caption("設定每個班級、每個科目「一週應該排幾節課」，系統會自動跟目前匯入的課表比對；下方也會列出教師被重複排課的衝突節次。")

        # 從 df_assign 重新解析出 assign_dict，確保節數檢查涵蓋「配課表有配、但課表完全沒排到」的科目
        _al, assign_dict_for_check, _tu, _at, _w = parse_assignment(st.session_state.df_assign)
        combos = sorted(collect_class_subject_combos(st.session_state.class_data, assign_dict_for_check))
        if "expected_hours" not in st.session_state:
            st.session_state.expected_hours = {}

        if not combos:
            st.info("目前沒有解析出任何科目，無法進行節數檢查。")
        else:
            st.subheader("① 設定應排節數")
            setting_rows = [
                {"班級": c, "科目": s, "應排節數": int(st.session_state.expected_hours.get(f"{c}|{s}", 0))}
                for c, s in combos
            ]
            setting_df = pd.DataFrame(setting_rows)

            edited_df = st.data_editor(
                setting_df,
                use_container_width=True,
                hide_index=True,
                key="expected_hours_editor",
                disabled=["班級", "科目"],
                column_config={
                    "應排節數": st.column_config.NumberColumn(
                        min_value=0, step=1, help="這個班級這個科目，一週應該排幾節課"
                    ),
                },
            )

            # 把編輯結果同步回 session_state，供雲端存檔與下方比對使用
            new_expected = {}
            for _, row in edited_df.iterrows():
                key = f"{row['班級']}|{row['科目']}"
                try:
                    new_expected[key] = int(row["應排節數"])
                except (ValueError, TypeError):
                    new_expected[key] = 0
            st.session_state.expected_hours = new_expected

            if st.button("💾 儲存節數設定"):
                ok, msg = github_save_data(build_save_payload())
                if ok:
                    st.success(f"☁️ {msg}")
                else:
                    st.warning(msg)

            st.divider()
            st.subheader("② 比對結果")

            actual_counts = compute_actual_subject_counts(st.session_state.class_data)
            check_rows = []
            for c, s in combos:
                expected = st.session_state.expected_hours.get(f"{c}|{s}", 0)
                actual = actual_counts.get((c, s), 0)
                diff = actual - expected
                if diff == 0:
                    status = "✅ 相符"
                elif diff > 0:
                    status = f"⚠️ 多排 {diff} 節"
                elif actual == 0:
                    status = f"❌ 完全沒排課（配課表有配 {expected} 節）"
                else:
                    status = f"❌ 少排 {abs(diff)} 節"
                check_rows.append({
                    "班級": c, "科目": s, "應排節數": expected, "實際節數": actual, "差異": diff, "狀態": status,
                })
            check_df = pd.DataFrame(check_rows)

            only_mismatch = st.checkbox("只顯示節數不相符的項目", value=True)
            display_df = check_df[check_df["差異"] != 0] if only_mismatch else check_df

            mismatch_count = int((check_df["差異"] != 0).sum())
            if mismatch_count == 0:
                st.success("🎉 所有已設定應排節數的科目，實際排課節數都相符！")
            else:
                st.warning(f"⚠️ 共有 {mismatch_count} 項班級科目的節數跟設定不相符")

            st.dataframe(
                display_df.sort_values(["班級", "科目"]).reset_index(drop=True),
                use_container_width=True, hide_index=True,
            )

        st.divider()
        st.subheader("③ 排課衝突偵測")
        st.caption("找出同一位教師在同一時段被排到多個不同班級的情況（通常是排課時的人為疏失）。")

        conflicts = detect_teacher_conflicts(st.session_state.class_data)
        if not conflicts:
            st.success("🎉 目前沒有偵測到任何排課衝突！")
        else:
            st.error(f"🚨 共偵測到 {len(conflicts)} 筆排課衝突")
            conflict_rows = [
                {
                    "教師": cf["teacher"],
                    "星期": f"週{cf['day']}",
                    "節次": f"第{cf['period']}節",
                    "同時被排到的班級": "、".join(cf["classes"]),
                }
                for cf in conflicts
            ]
            st.dataframe(pd.DataFrame(conflict_rows), use_container_width=True, hide_index=True)

else:
    st.info("👋 請上傳資料檔並點擊執行整合。")
