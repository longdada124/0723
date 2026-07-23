import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re
import os
import unicodedata

st.set_page_config(page_title="課表彙整系統", layout="wide")

# --- 核心替換函數 ---
def master_replace(doc_obj, old_text, new_text):
    if isinstance(new_text, (float, int)):
        new_val = str(int(new_text))
    else:
        new_val = str(new_text) if (new_text and str(new_text).strip() != "") else ""
    targets = list(doc_obj.paragraphs)
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                targets.extend(cell.paragraphs)
    for p in targets:
        if old_text in p.text:
            full_text = "".join([run.text for run in p.runs])
            updated_text = full_text.replace(old_text, new_val)
            for i, run in enumerate(p.runs):
                run.text = updated_text if i == 0 else ""

# --- 讀取內建樣板函數 (編碼相容版) ---
def load_default_template(file_name):
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        full_path = os.path.join(current_dir, file_name)
        
        if os.path.exists(full_path):
            with open(full_path, "rb") as f:
                return f.read()
        
        target_filename_nfc = unicodedata.normalize('NFC', file_name)
        for actual_file in os.listdir(current_dir):
            actual_file_nfc = unicodedata.normalize('NFC', actual_file)
            if actual_file_nfc == target_filename_nfc:
                with open(os.path.join(current_dir, actual_file), "rb") as f:
                    return f.read()
    except Exception:
        pass
    return None

# --- 側邊欄：資料管理 ---
with st.sidebar:
    st.header("⚙️ 資料管理")
    if st.button("🧹 清空重置系統"):
        for key in list(st.session_state.keys()): del st.session_state[key]
        st.rerun()

    st.divider()
    st.subheader("📥 範本下載")
    data_templates = {
        "1. 配課表範本": "配課表.xlsx",
        "2. 課表範本": "課表.xlsx",
        "3. 教師排序表範本": "教師排序表.xlsx"
    }
    for label, file_name in data_templates.items():
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            f_path = os.path.join(current_dir, file_name)
            if os.path.exists(f_path):
                with open(f_path, "rb") as f:
                    st.download_button(label=f"{label}", data=f, file_name=file_name, key=f"dl_{file_name}")
            else:
                st.caption(f"⚠️ 找不到 {file_name}")
        except Exception:
            st.caption(f"⚠️ 讀取 {file_name} 失敗")
    st.divider()

    st.subheader("📤 上傳資料檔")
    f_assign = st.file_uploader("1. 上傳【配課表】", type=["xlsx", "csv"])
    f_time = st.file_uploader("2. 上傳【課表】", type=["xlsx", "csv"])
    f_sort = st.file_uploader("3. 上傳【教師排序暨時數表】", type=["xlsx", "csv"])
    
    if f_assign and f_time and st.button("🚀 執行整合"):
        class_temp = load_default_template("班級樣板.docx")
        teacher_temp = load_default_template("教師樣板.docx")
        
        if not class_temp or not teacher_temp:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            actual_files = os.listdir(current_dir) if os.path.exists(current_dir) else []
            st.error("❌ 系統錯誤：後台找不到「班級樣板.docx」或「教師樣板.docx」，請確認 GitHub 檔案。")
            st.info(f"🔍 雲端伺服器目前實際看到的檔案清單：\n`{actual_files}`")
        else:
            with st.spinner("同步內建樣板與解析資料中..."):
                df_assign = pd.read_csv(f_assign) if f_assign.name.endswith('.csv') else pd.read_excel(f_assign)
                df_time = pd.read_csv(f_time) if f_time.name.endswith('.csv') else pd.read_excel(f_time)
                
                st.session_state.class_template = class_temp
                st.session_state.teacher_template = teacher_temp
                st.session_state.df_assign = df_assign

                # --- 1. 解析配課 (智慧雙模解析) ---
                assign_lookup = []
                all_teachers_db = set()
                tutors = {}

                df_assign.columns = [str(c).strip() for c in df_assign.columns]

                if '科目' in df_assign.columns and '教師' in df_assign.columns:
                    for _, row in df_assign.iterrows():
                        c, s, t_raw = str(row['班級']).strip(), str(row['科目']).strip(), str(row['教師']).strip()
                        t_list = [name.strip() for name in t_raw.split('/')]
                        for t in t_list:
                            if t and t != "nan":
                                assign_lookup.append({'c': c, 's': s, 't': t})
                                all_teachers_db.add(t)
                        if s == "班級": 
                            tutors[c] = t_raw
                else:
                    if '導師' in df_assign.columns:
                        for _, row in df_assign.iterrows():
                            c = str(row['班級']).strip()
                            t_tutor = str(row['導師']).strip()
                            if t_tutor and t_tutor != "nan":
                                tutors[c] = t_tutor

                    subject_cols = [col for col in df_assign.columns if col not in ['班級', '導師']]
                    df_melted = pd.melt(df_assign, id_vars=['班級'], value_vars=subject_cols, var_name='科目', value_name='教師')
                    
                    for _, row in df_melted.iterrows():
                        c = str(row['班級']).strip()
                        s = str(row['科目']).strip()
                        t_raw = str(row['教師']).strip() if pd.notna(row['教師']) else ""
                        if not t_raw or t_raw == "nan" or t_raw == "": continue
                        
                        t_list = [name.strip() for name in t_raw.split('/')]
                        for t in t_list:
                            if t and t != "nan" and t != "":
                                assign_lookup.append({'c': c, 's': s, 't': t})
                                all_teachers_db.add(t)

                # --- 2. 教師排序與時數 ---
                ordered_teachers, base_hours, all_teachers_list = [], {}, list(all_teachers_db)
                if f_sort:
                    df_s = pd.read_csv(f_sort) if f_sort.name.endswith('.csv') else pd.read_excel(f_sort)
                    for _, s_row in df_s.iterrows():
                        t_name = str(s_row.iloc[0]).strip()
                        if t_name in all_teachers_list:
                            ordered_teachers.append(t_name)
                            try: base_hours[t_name] = int(s_row.iloc[1])
                            except: base_hours[t_name] = 0
                    for t in all_teachers_list:
                        if t not in ordered_teachers: ordered_teachers.append(t); base_hours[t] = 0
                else:
                    ordered_teachers = sorted(all_teachers_list)
                    base_hours = {t: 0 for t in ordered_teachers}

                # --- 3. 解析課表 (全新雙模解析：一維清單 & 二維矩陣) ---
                parsed_timetable = []
                df_time.columns = [str(c).strip() for c in df_time.columns]

                # 【模式 A】舊版一維清單 (有星期、節次欄位)
                if '星期' in df_time.columns and '節次' in df_time.columns:
                    day_map = {"一":1,"二":2,"三":3,"四":4,"五":5,"週一":1,"週二":2,"週三":3,"週四":4,"週五":5}
                    for _, row in df_time.iterrows():
                        c_raw, s_raw = str(row['班級']).strip(), str(row['科目']).strip()
                        if not s_raw or s_raw == "nan" or s_raw == "": continue
                        d = day_map.get(str(row['星期']).strip(), 0)
                        p_match = re.search(r'\d+', str(row['節次']))
                        if p_match and d > 0:
                            parsed_timetable.append({'c': c_raw, 'd': d, 'p': int(p_match.group()), 's': s_raw})
                
                # 【模式 B】新版二維極簡矩陣表 (只有班級 + 各時段如 一1, 週二3)
                else:
                    # 找出表示時段的欄位名稱
                    time_cols = [col for col in df_time.columns if col != '班級']
                    # 將二維表格融化為一維清單，方便後續處理
                    df_time_melted = pd.melt(df_time, id_vars=['班級'], value_vars=time_cols, var_name='時段', value_name='科目')
                    
                    day_map = {"一":1, "二":2, "三":3, "四":4, "五":5}
                    for _, row in df_time_melted.iterrows():
                        c_raw, s_raw = str(row['班級']).strip(), str(row['科目']).strip()
                        if not s_raw or s_raw == "nan" or s_raw == "": continue
                        
                        # 從字串中萃取星期與節次 (例如 "一1" 或 "週三4")
                        t_str = str(row['時段']).strip().replace('週', '')
                        d_match = re.search(r'(一|二|三|四|五)', t_str)
                        p_match = re.search(r'\d+', t_str)
                        
                        if d_match and p_match:
                            d = day_map.get(d_match.group(1), 0)
                            p = int(p_match.group())
                            parsed_timetable.append({'c': c_raw, 'd': d, 'p': p, 's': s_raw})

                # 將標準化後的排課資料組合起來
                class_data, teacher_data, total_counts = {}, {}, {}
                for item in parsed_timetable:
                    c_raw, d, p, s_raw = item['c'], item['d'], item['p'], item['s']
                    
                    # 比對配課表，找出授課教師
                    #curr_t_list = [a檢視了您目前的程式碼[cite: 3] 以及上傳的 `課表.xlsx` 結構後，發現目前的課表採用的是「一維資料庫」的清單格式（即每一列包含 `班級`、`星期`、`節次`、`科目`）。這代表若一個班級一週有 35 節課，光是一個班就要填寫 35 列，全校若有 30 個班，檔案將會長達 1000 多列，這對行政人員填寫來說確實非常複雜且難以閱讀。

### 💡 建議的全新 `課表.xlsx` 超極簡寫法

您可以將 Excel 改為一般人最習慣閱讀的「二維橫表」，**每一列就是一個班級，後面的欄位直接對應「星期幾的第幾節」**。
像下方這樣，只要短短幾列，就能清楚列出全校的課表：

| 班級 | 一1 | 一2 | 一3 | 一4 | ... | 五7 | 五8 |
| :--- | :--- | :--- | :--- | :--- | :--- | :--- | :--- |
| **901** | 全球化 | 自然 | 國文 | 體育 | ... | 班會 | 社團 |
| **902** | 英文 | 國文 | 歷史 | 數學 | ... | 班會 | 輔導 |

*(備註：上方的欄位名稱您寫 `一1` 或是 `週一1` 都可以，程式都會自動精準辨識！空白不填的格子程式就會當作那一節沒有排課。)*

---

### 🚀 升級版程式碼（自動支援「一維舊版」與「二維極簡版」雙模課表）

我已經將您的 `# --- 3. 解析課表 ---` 區塊改寫為**智慧雙模架構**。不論您上傳的是舊版的冗長清單，還是上述全新的二維極簡表，系統都會自動判斷並處理。

請將您原程式碼中的 `# --- 3. 解析課表 ---` 區塊（大約在 176 行處）替換為以下程式碼：

```python
                # --- 3. 解析課表 (支援智慧三模解析：一維傳統表、橫向班級表、直向時段表) ---
                class_data, teacher_data, total_counts = {}, {}, {}
                day_map = {"一":1,"二":2,"三":3,"四":4,"五":5,"週一":1,"週二":2,"週三":3,"週四":4,"週五":5}
                
                # 確保欄位名稱乾淨
                df_time.columns = [str(c).strip() for c in df_time.columns]
                
                parsed_time_list = []
                
                if '星期' in df_time.columns and '節次' in df_time.columns and '科目' in df_time.columns:
                    # 【模式 A】舊版一維清單格式 (班級, 星期, 節次, 科目)
                    for _, row in df_time.iterrows():
                        c_raw, s_raw = str(row['班級']).strip(), str(row['科目']).strip()
                        d = day_map.get(str(row['星期']).strip(), 0)
                        p_match = re.search(r'\d+', str(row['節次']))
                        
                        if p_match and d > 0 and s_raw and s_raw != "nan" and s_raw != "":
                            parsed_time_list.append({'c': c_raw, 's': s_raw, 'd': d, 'p': int(p_match.group())})
                else:
                    # 【模式 B】二維矩陣格式 (自動判定是「橫向」還是「直向」)
                    if '班級' in df_time.columns:
                        # B1: 橫向表 (班級在列，欄位是時段如「一1」)
                        id_col = '班級'
                        val_cols = [col for col in df_time.columns if col != '班級']
                        var_name, val_name = '時段', '科目'
                    else:
                        # B2: 直向表 (時段在列，欄位是班級如「901」) 👈 你最新提議的最佳解！
                        id_col = df_time.columns[0] # 將第一欄視為時段 (例如欄位名為「節次」)
                        val_cols = [col for col in df_time.columns if col != id_col]
                        var_name, val_name = '班級', '科目'

                    # 利用 pandas melt 將二維表格融化為一維清單
                    df_time_melted = pd.melt(df_time, id_vars=[id_col], value_vars=val_cols, var_name=var_name, value_name=val_name)
                    
                    for _, row in df_time_melted.iterrows():
                        # 動態取得班級與時段的值
                        c_raw = str(row['班級']).strip() if var_name == '班級' else str(row[id_col]).strip()
                        s_raw = str(row['科目']).strip()
                        t_raw = str(row[id_col]).strip() if var_name == '班級' else str(row[var_name]).strip()
                        
                        if not s_raw or s_raw == "nan" or s_raw == "": continue
                        
                        # 萃取星期與節次
                        t_str = t_raw.replace('週', '')
                        d_match = re.search(r'(一|二|三|四|五)', t_str)
                        p_match = re.search(r'\d+', t_str)
                        
                        if d_match and p_match:
                            d = day_map.get(d_match.group(1), 0)
                            p = int(p_match.group())
                            parsed_time_list.append({'c': c_raw, 's': s_raw, 'd': d, 'p': p})

                # 將標準化後的排課資料組合起來
                for item in parsed_time_list:
                    c_raw, s_raw, d, p = item['c'], item['s'], item['d'], item['p']
                    
                    # 從配課表查找授課教師
                    curr_t_list = [a['t'] for a in assign_lookup if a['c'] == c_raw and a['s'] == s_raw]
                    display_t = "/".join(curr_t_list) if curr_t_list else "未知教師"
                    
                    if c_raw not in class_data: class_data[c_raw] = {}
                    class_data[c_raw][(d, p)] = {"subj": s_raw, "teacher": display_t}
                    
                    for t in curr_t_list:
                        if t not in teacher_data: teacher_data[t] = {}
                        teacher_data[t][(d, p)] = {"subj": s_raw, "class": c_raw}
                        total_counts[t] = total_counts.get(t, 0) + 1

                st.session_state.update({
                    "class_data": class_data, "teacher_data": teacher_data, "tutors_map": tutors,
                    "base_hours": base_hours, "total_counts": total_counts, "ordered_teachers": ordered_teachers,
                    "sel_class": sorted(list(class_data.keys()))[0] if class_data else "", 
                    "sel_teacher": ordered_teachers[0] if ordered_teachers else ""
                })
                st.rerun()
